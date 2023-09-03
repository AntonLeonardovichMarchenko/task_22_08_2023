import os
import shutil
import zipfile

fileNameSource = "C:\\PythonDrom\\Tests_2022\\test_22_08_2023\\ANSI_ShipsBook.txt"

#Импорт необходимых библиотек

from docxtpl import DocxTemplate
from docxtpl import InlineImage

from docxtpl import DocxTemplate, InlineImage

import fileinput, glob, os

from docx import Document
from docx.shared import Inches, Cm, Mm, Pt


class fileMaster:

    fm=None


    def __init__(self, file, mode):

        if self.fm != None:
            self.fm.close()

        self.fm = open(file,
                      mode,
                      buffering=-1,
                      encoding=None,
                      errors=None,
                      newline=None,
                      closefd=True,
                      opener=None)

    def write_FM(self, strKey):
        self.fm.write(strKey)

    def read_FM(self):
        return self.fm.readline()



class WarShip:
    NumberOfShip = 0
    classOfShip = ''
    descriptionOfShip = ''
    TheWarShips = list()

    fm = fileMaster(fileNameSource, 'r')

    # создание записи о ЦЕЛОЙ категории судов (всё за один раз!)
    # номер-категория_судна-описание_категории
    # Так как ВСЕ записи имеют общий формат, это легко.
    def __init__(self):
        # из файла fm в переменную ws с помощью метода WarShip.fm.read_FM()
        # в локальную переменную читается строка ws, которая содержит информацию
        # о категории судна...
        ws = WarShip.fm.read_FM()
        # значение ws записывается в поле self.classOfShip объекта self.classOfShip
        self.classOfShip = ws
        # в поле self.numberOfShip записывается текущее значение переменной класса,
        # которая автоматически увеличивается после заполнения очередной записи о
        # категории судов
        self.numberOfShip = WarShip.NumberOfShip


        # чтение многострочного описания_категории судна. Признаком конца очередного
        # описания_категории судна является пустая символ-строка перехода на следующую
        # строку в файле WarShip.fm
        cf = ''
        while cf!='\n':
            cf = WarShip.fm.read_FM()    # строка описания читается в cf
            self.descriptionOfShip = self.descriptionOfShip + cf        # cf добавляется в поле
                                                                        # self.descriptionOfShip
                                                                        # объекта - представителя
                                                                        # класса WarShip


        WarShip.NumberOfShip = WarShip.NumberOfShip+1    # увеличивается значение переменной
                                                         # класса WarShip.NumberOfShip



#
# def GenerateReport(docxName,
#                    reportName,
#                    head,
#                    fleet,
#                    listShips,
#                    numShip,
#                    classShip,
#                    descriptShip):
#
#     signature = 'DreadNought_F.png'
#     FromTemplate(
#                  docxName,
#                  reportName,
#                  signature,
#                  head,
#                  fleet,
#                  listShips,
#                  numShip,
#                  classShip,
#                  descriptShip,
#                  )
#
# def FromTemplate(
#                  docxName,
#                  reportName,
#                  signature,
#                  head,
#                  fleet,
#                  listShips,
#                  numShip,
#                  classShip,
#                  descriptShip,
#                  ):
#
#
# # ========================================================================
#
#     template = DocxTemplate(docxName)
#     img_size = Cm(15)
#     acc = InlineImage(template, signature, img_size) # для вставки картинки
#
#
#     # в функцию передаются только значения.
#     # в данном случае это вообще ни с чем
#     # не связанные строковые значения.
#     #                                   Соответствующие ключи забиты в функцию
#     # context = GetContext(acc,                   # accPosition
#     #                      'great Fleet',         # headLine
#     #                      '~~~ the fleet ~~~',   # greateFleet
#     #                      'list of Ships',       # listOfShips
#     #                      'num of Ship',         # numberOfShips
#     #                      'class of Ship',       # classOfShips
#     #                      'description of Ship'  # descriptionOfShips
#     #                      )   # объекты словаря
#
#     # здесь словарь context заполнен значениями - аргументами функции FromTemplate
#     # context = GetContext(acc,           # accPosition
#     #                      head,          # headLine
#     #                      fleet,         # greateFleet
#     #                      listShips,     # listOfShips
#     #                      numShip,       # numberOfShips
#     #                      classShip,     # classOfShips
#     #                      descriptShip   # descriptionOfShips
#     #                      )   # context содержит объекты словаря
#
#     # таким образом, удаётся автоматизировать формирование словаря (контекста),
#     # но возникает проблема формирования списков аргументов.
#     keys = ['accPosition', 'headLine', 'greateFleet', 'listOfShips', 'numberOfShips', 'classOfShips', 'descriptionOfShips']
#     values = [acc, head, fleet, listShips, numShip, classShip, descriptShip]
#
#     context = GetContext_2(keys, values)
#
#
#     template.render(context)
#
#
#     # можно соответствующий словарь создать непосредственно в выражении вызова метода
#     # template.render. Так даже нагляднее.
#     # template.render({'accPosition': acc,
#     #               'headLine': head,
#     #               'greateFleet': fleet,
#     #               'listOfShips': listShips,
#     #               'numberOfShips': numShip,
#     #               'classOfShips': classShip,
#     #               'descriptionOfShips': descriptShip})
#
#
#
#     template.save(reportName)
#
#     # ========================================================================
#
#     #template_2 = DocxTemplate("word_tmpl.docx")  # шаблон
#     #img_size = Cm(15)
#     # acc = InlineImage(template_2, signature, img_size) # для вставки картинки
#     #
#     #
#     # # в функцию передаются только значения.
#     # # в данном случае это вообще ни с чем
#     # # не связанные строковые значения.
#     # #                                   Соответствующие ключи забиты в функцию
#     # # context = GetContext(acc,                   # accPosition
#     # #                      'great Fleet',         # headLine
#     # #                      '~~~ the fleet ~~~',   # greateFleet
#     # #                      'list of Ships',       # listOfShips
#     # #                      'num of Ship',         # numberOfShips
#     # #                      'class of Ship',       # classOfShips
#     # #                      'description of Ship'  # descriptionOfShips
#     # #                      )   # объекты словаря
#     #
#     # # здесь словарь context заполнен значениями - аргументами функции FromTemplate
#     # # context = GetContext(acc,           # accPosition
#     # #                      head,          # headLine
#     # #                      fleet,         # greateFleet
#     # #                      listShips,     # listOfShips
#     # #                      numShip,       # numberOfShips
#     # #                      classShip,     # classOfShips
#     # #                      descriptShip   # descriptionOfShips
#     # #                      )   # context содержит объекты словаря
#     #
#     # # таким образом, удаётся автоматизировать формирование словаря (контекста),
#     # # но возникает проблема формирования списков аргументов.
#     # keys = ['accPosition', 'headLine', 'greateFleet', 'listOfShips', 'numberOfShips', 'classOfShips', 'descriptionOfShips']
#     # values = [acc, head, fleet, listShips, numShip, classShip, descriptShip]
#     #
#     # context = GetContext_2(keys, values)
#     #
#     #
#     # template_2.render(context)
#     #
#     #
#     # # можно соответствующий словарь создать непосредственно в выражении вызова метода
#     # # template.render. Так даже нагляднее.
#     # # template.render({'accPosition': acc,
#     # #               'headLine': head,
#     # #               'greateFleet': fleet,
#     # #               'listOfShips': listShips,
#     # #               'numberOfShips': numShip,
#     # #               'classOfShips': classShip,
#     # #               'descriptionOfShips': descriptShip})
#     #
#     # template_2.save('report.docx')
#
#
#
# # в этой версии функции имена полей (места вставок значений в шаблоне документа) забиты
# # непосредственно в тело функции.
# # В принципе, можно было бы попробовать перед обращением к подобной функции сформировать
# # ДВА списка:
# # список мест подстановки в документе и
# # список значений для подстановки.
# # Пока только ПОКА не понятно, как организовать переменное количество полей подстановки
# # непосредственно на странице документа.
# def GetContext(acc, theHead, theFleet, theList, theNum, theClass, theDescription):
#     return {
#             'accPosition':        acc,
#             'headLine':           theHead,
#             'greateFleet':        theFleet,
#             'listOfShips':        theList,
#             'numberOfShips':      theNum,
#             'classOfShips':       theClass,
#             'descriptionOfShips': theDescription
#     }       # построение и возвращение словаря
#
#
# def GetContext_2(keys, values):
#
#     context = dict()
#     for i in range(len(keys)):
#         context[keys[i]] = values[i]
#
#     return context
#
#
#
#
def DoIt():

    fm = None
#   ws = ''
    #fm = open("report_fm.docx", "wb+")

      # список описаний классов судов. Чтение информации о классе судна.
    while True:
        ws = WarShip()
        if ws.classOfShip != '~~~~~~~~~~:' and ws.descriptionOfShip != '\n':
            # читать из файла пока не будет прочитана эта строка
            # (признак конца файла описания)
            WarShip.TheWarShips.append(ws)
        else:
            break

    # файл с описаниями прочитан, информация собрана в записи-представители класса WarShip и
    # собраны в список TheWarShips (смотреть функцию DoIt)

    # всего лишь распечатка результатов
    for ws in WarShip.TheWarShips:
        print(ws.numberOfShip, ws.classOfShip, ws.descriptionOfShip)

    # ====================================================================


     # tpl = DocxTemplate('fleet_tpl_X.docx')
     #
     # # python-docx is a Python library for creating and updating Microsoft Word (.docx) files.
     #
     # document = Document()
     # document.add_heading('Document Title', 0)
     #
     # p = document.add_paragraph('A plain paragraph having some ')
     # p.add_run('bold').bold = True
     # p.add_run(' and some ')
     # p.add_run('italic.').italic = True
     #
     # document.add_heading('Heading, level 1', level=1)
     # document.add_paragraph('Intense quote', style='Intense Quote')
     #
     # document.add_paragraph(
     #     'first item in unordered list', style='List Bullet'
     # )
     # document.add_paragraph(
     #     'first item in ordered list', style='List Number'
     # )
     #
     # document.add_picture('DreadNought_F.png', width=Inches(1.25))
     #
     #
     #
     # records = (
     #     (WarShip.TheWarShips[2].numberOfShip, WarShip.TheWarShips[2].classOfShip, WarShip.TheWarShips[2].descriptionOfShip),
     #     (WarShip.TheWarShips[0].numberOfShip, WarShip.TheWarShips[0].classOfShip, WarShip.TheWarShips[0].descriptionOfShip),
     #     (WarShip.TheWarShips[5].numberOfShip, WarShip.TheWarShips[5].classOfShip, WarShip.TheWarShips[5].descriptionOfShip)
     # )
     #
     # table = document.add_table(rows=1, cols=3)
     # hdr_cells = table.rows[0].cells
     # hdr_cells[0].text = 'Qty'
     # hdr_cells[1].text = 'Id'
     # hdr_cells[2].text = 'Desc'
     # # for qty, id, desc in records:
     # #     row_cells = table.add_row().cells
     # #     row_cells[0].text = str(qty)
     # #     row_cells[1].text = id
     # #     row_cells[2].text = desc
     #
     # document.add_page_break()
     #
     # document.save('fleetText.docx')
     #
     # # ====================================================================
     #
     # source_file = "word_tmpl.docx"
     # destination_file = "word_tmpl_1.docx"


     # GenerateReport(
     #                "word_tmpl_1.docx",
     #                "report_1.docx",
     #                '>>> 1 the greate fleet',
     #                'The Fleet',
     #                'FFF',
     #                'list Of ships',
     #                'asdfg',
     #                'qqqq'
     #                )

     #fm = open("report_fm.docx", "wb+")
     #fs = open("report_1.docx", 'rb')
     #shutil.copyfileobj(fs, fm)
     #fm.write('xxx'.encode())

     # for b in fs:
     #     #fm.write(b)
     #     #fm.write(str(b).encode())
     #     fm.write('1'.encode())
     #     # fm.write('\n'.encode())
     #     fm.write('2'.encode())
     #
     #     #fm.writelines(['\n'.encode(),'\n'.encode(), '\n'.encode(), '\n'.encode()])
     #     fs = open("report_1.docx", 'rb')
     #
     #     #print(b)

     # fm.write('\n'.encode())
     # fm.write('1'.encode())



     # for line in fs:
     #     print(f"{line}")

     #fm.fm.write(fs)
     # cf = ''
     # while cf != '\n':
     #     cf = fs.readline()
     # #     fm.fm.write(cf)

     #fm.write_FM(f"  {fs}")

     # source_file = "word_tmpl.docx"
     # destination_file = "word_tmpl_2.docx"
     # shutil.copy2(source_file, destination_file)

     # GenerateReport(
     #     "word_tmpl_2.docx",
     #     "report_2.docx",
     #     '>>> 2 the greate fleet',
     #     'The Fleet',
     #     'GGG',
     #     'list Of ships',
     #     'sdfgh',
     #     'wwww'
     # )

     # #fm.write('\n\n\n\nasdfg\n\n\n'.encode())
     #
     # hi = 'привет'
     # he = hi.encode('utf-16')
     # #hb = b'\xd0\xbf\xd1\x80\xd0\xb8\xd0\xb2\xd0\xb5\xd1\x82\xd0\xbf\xd1\x80\xd0\xb8\xd0\xb2\xd0\xb5\xd1\x82'
     #
     #
     # print(he)
     # #print(hb)
     # print(he.decode('utf-16'))
     # #print(hb.decode('utf-16'))
     #
     # print('~~~~~~~~~~~~~~~~~~~~~~~~~~')
     # hi = '\n\n\n\n{QWeee}'
     # he = hi.encode('utf-16')
     #
     # print(he)
     # print(he.decode('utf-16'))
     #
     #
     #
     #
     # fs = open("report_2.docx", 'rb')
     #
     # for b in fs:
     # #     fm.write(b)
     #     print(b)
     #     fm.write(b)
     #
     # fm.write(he)

     # listFiles = []
     # listFiles.append(open('report_1.docx', 'rb'))
     # listFiles.append(open('report_2.docx', 'rb'))
     # fw = open('report_fm.docx', 'w')
     #
     # for i in range(0,2):
     #
     #     #fw.write('123456789'.encode('utf-16'))
     #
     #     print(listFiles[i])
     #
     #     for line in listFiles[i]:
     #         print(line)
     #         fw.write(str(line))
     #         #fw.write('qwerty'.encode('utf-16'))
     #         print(f'++++++++++++++++++++++++++{i}++++++++++++++++++++++++++++++++++++')

     #
     #     line = 'wwwwwwwwwwwwwwwwwwwwwwwwwwwwwwwwwwwwwwwwwwwww\n\n\n123'.encode('utf-16')
     #     print (line)
     #     fw.write(line)
     #     fw.close()

    print("**********************************************************************")
    document = Document()
    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.add_picture('DreadNought_F.png', width=Inches(1.25))





    records = (
        (WarShip.TheWarShips[1].numberOfShip, WarShip.TheWarShips[1].classOfShip, WarShip.TheWarShips[1].descriptionOfShip),
        (WarShip.TheWarShips[0].numberOfShip, WarShip.TheWarShips[1].classOfShip, WarShip.TheWarShips[0].descriptionOfShip),
        (WarShip.TheWarShips[2].numberOfShip, WarShip.TheWarShips[2].classOfShip, WarShip.TheWarShips[2].descriptionOfShip)
    )

    print(WarShip.TheWarShips[0].classOfShip)

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'numberOfShip'
    hdr_cells[1].text = 'classOfShip'
    hdr_cells[2].text = 'descriptionOfShip'

    for _numberOfShip, _classOfShip, _descriptionOfShip in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(_numberOfShip)
        row_cells[1].text = _classOfShip
        row_cells[2].text = _descriptionOfShip

    document.add_page_break()
    document.save('mainTest.docx')



    ##document.add_heading('Heading, level 1', level=1)
   ##document.add_paragraph('Intense quote', style='Intense Quote')
    ##document.add_paragraph(
   ##    'first item in unordered list', style='List Bullet'
   ##)
   ##document.add_paragraph(
   ##    'first item in ordered list', style='List Number'
   ##)
    ##document.add_picture('DreadNought_F.png', width=Inches(1.25))

   ##records = (
   ##    Ship.TheWarShips[0].numberOfShip, WarShip.TheWarShips[0].classOfShip,WarShip.TheWarShips[0].descriptionOfShip)
    ##)
   ##for ws in WarShip.TheWarShips:
   ##    print(ws.numberOfShip, ws.classOfShip, ws.descriptionOfShip)



   ##records = (
    ##    (3, '101', 'Spam'),
    ##    (7, '422', 'Eggs'),
    ##    (4, '631', 'Spam, spam, eggs, and spam')
    ##)

    # context = {}
    # context['the fleet'] = 'Это о кораблях. Больших и маленьких, новых и старых.'
    # doc = DocxTemplate('report_1.docx')
    # doc.render(context)
    # doc.save('report_full.docx')






def main():
    DoIt()



if __name__ == "__main__":
    main()
