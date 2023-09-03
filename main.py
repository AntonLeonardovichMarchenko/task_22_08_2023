import time
import csv
import json

fileNameSource = "C:\\PythonDrom\\Tests_2022\\test_22_08_2023\\ANSI_ShipsBook.txt"

from docxtpl import DocxTemplate, InlineImage, RichText
from docx import Document
from docx.shared import Inches, Cm, Mm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

# ========================================================================
#      класс fileMaster отвечает за открытие-закрытие и чтение-запись

class fileMaster:

    fm=None

    def __init__(self, file, mode):

        if self.fm != None:
            self.fm.close()

        self.fm = open(file,
                      mode,
                      buffering=-1,
                      encoding=None,
                      errors=None,
                      newline=None,
                      closefd=True,
                      opener=None)

    def write_FM(self, strKey):
        self.fm.write(strKey)

    def read_FM(self):
        return self.fm.readline()
# ========================================================================


class universalFileMaster:

    fm = None

    def __init__(self, file, mode):


        if universalFileMaster.fm is not None:
            universalFileMaster.fm.close()

        universalFileMaster.fm = open(file,
                            mode,
                            buffering=-1,
                            encoding=None,
                            errors=None,
                            newline=None,
                            closefd=True,
                            opener=None)

# разница между классами fileMaster и universalFileMaster заключается
# в том, что в fileMaster ссылка на файл является локальной, становится
# атрибутом конкретного класса (в данном случае класса WarShip) и для неё
# определены методы чтения-записи, а в universalFileMaster ссылка на файл
# является глобальной и методы чтения-записи определяются в классах,
# реализующих конкретные технологии вводе-вывода. Хотя тут, может быть,
# есть недоработка при реализации принципов ООП.
# ========================================================================


class WarShip:
    NumberOfShip = 0
    classOfShip = ''
    descriptionOfShip = ''
    TheWarShips = list()

    fm = fileMaster(fileNameSource, 'r')

    # создание записи о ЦЕЛОЙ категории судов (всё за один раз!)
    # номер-категория_судна-описание_категории
    # Так как ВСЕ записи имеют общий формат, это легко.
    def __init__(self):
        # из файла fm в переменную ws с помощью метода WarShip.fm.read_FM()
        # в локальную переменную читается строка ws, которая содержит информацию
        # о категории судна...
        ws = WarShip.fm.read_FM()
        # значение ws записывается в поле self.classOfShip объекта self.classOfShip
        self.classOfShip = ws
        # в поле self.numberOfShip записывается текущее значение переменной класса,
        # которая автоматически увеличивается после заполнения очередной записи о
        # категории судов
        self.numberOfShip = WarShip.NumberOfShip


        # чтение многострочного описания_категории судна. Признаком конца очередного
        # описания_категории судна является пустая символ-строка перехода на следующую
        # строку в файле WarShip.fm
        cf = ''
        while cf!='\n':
            cf = WarShip.fm.read_FM()    # строка описания читается в cf
            self.descriptionOfShip = self.descriptionOfShip + cf        # cf добавляется в поле
                                                                        # self.descriptionOfShip
                                                                        # объекта - представителя
                                                                        # класса WarShip


        WarShip.NumberOfShip = WarShip.NumberOfShip+1    # увеличивается значение переменной
                                                         # класса WarShip.NumberOfShip

# ========================================================================
# обмен данными. эти данные бывают разных видов, например файл, строка или число.
# Есть структура данных, которая позволяет быстро и просто воссоздавать объекты (???) и
# обмениваться этими данными по сети, — это JSON
# JSON — это строка со словарем. Она представлена в виде байтовой последовательности.
# Можно отправить ее по сети приложению, а в нём воссоздать полученную структуру
# в объекты языка (???).
#
# Сериализация и десериализация
# В Python есть множество библиотек, чтобы работать с JSON, в том числе и встроенная
# библиотека JSON Python. Она позволяет приводить любые структуры данных к JSON-объекту,
# — вплоть до пользовательских классов.
# А из этого объекта получать совместимую для работы в Python сущность — объект языка.
#
# Упаковка объектов в байтовую последовательность называется сериализацией.
# Распаковка байтов в объекты языка программирования, приведение последовательности
# назад к типам и структурам, — десериализацией.
# В байты данные необходимо переводить, чтобы отправлять их по сети или локально другому
# приложению, так как иной формат передать невозможно.
#
# Сериализация JSON
# Сериализация – это метод преобразования объектов Python в JSON.
# Иногда компьютеру требуется обработать большой объем информации, поэтому рекомендуется
# сохранить эту информацию в файле. Данные JSON можно сохранить в файле с помощью функции
# JSON. Для эого в модуле json есть методы dump() и dumps(), которые используются
# для преобразования объекта Python. Это функции для передачи (кодирования) данных в формате
# JSON. dump осуществляет запись данных JSON в файл. Она принимает два позиционных аргумента:
# первый – это объект данных, который нужно сериализовать, а второй – файловый объект,
# в который должны быть записаны байты.
#
# Десериализация JSON
# Десериализация – это процесс декодирования данных JSON в объекты Python.
# Модуль json предоставляет два метода load() и loads(), которые используются для
# преобразования данных JSON в фактическую объектную форму Python. Список соответствия
# приведен ниже:
#
#       JSON	Python
# 1.	Object	Dict
# 2.	Array	list
# 3.	String	str
# 4.	Number (int)	int
# 5.	true	True
# 6.	false	False
# 7.	null	None
#
# В приведенной таблице показано обратное преобразование сериализованной таблицы.
# но технически это не идеальное преобразование данных JSON. Это означает, что если
# объект кодируется и снова декодирутся, то это будет ДРУГОЙ объект.
#
# Реальный пример: один человек переводит что-то с английского языка
# на китайский язык, а другой переводит обратно (на английский), это может быть
# переведено неточно.
#
# Пример
#
# import json
# a = (10,20,30,40,50,60,70)
# print(type(a))
# b = json.dumps(a)
# print(type(json.loads(b)))
#
# Выход:
#
# <class 'tuple'>
# <class 'list'>
#
#
# Функции
# dumps позволяет создать JSON-строку из переданного в нее объекта.
# loads — преобразовать строку назад в объекты языка.
#
# dump чтобы сохранить результат в файл
# load используют для воссоздания объекта.
#
# Работают они схожим образом, но dump и load требуют передачи специального
# объекта для работы с файлом — filehandler.
#
# Работа с пользовательскими объектами
# Пользовательские классы (в том числе, WarShip) не относятся к JSON-сериализуемым.
# Это значит, что просто применить к ним функции dumps, loads или dump и load не
# получится.
# Один из вариантов решения проблемы - паттерн "Адаптер".
# Это класс, класс, который приводит к JSON пользовательские объекты и восстанавливает их.
#
# ...

# потом объявлю класс и работу с json переделаю
# class jsonMaster:
#     jsonFM = None
#     def __init__(self, fileNameDestination):
#         jsonMaster.jsonFM = fileMaster(fileNameDestination, 'w')
#
#     def jsonWrite(self, dataList):
#         with jsonMaster.jsonFM:
#             writer = json.writer(jsonMaster.jsonFM)
#             writer.writerows()

# ========================================================================

def DoIt():
    # record start time
    start = time.time()

    # список описаний классов судов. Чтение информации о классе судна.
    while True:
        ws = WarShip()
        if ws.classOfShip != '~~~~~~~~~~:' and ws.descriptionOfShip != '\n':
            # читать из файла пока не будет прочитана эта строка
            # (признак конца файла описания)
            WarShip.TheWarShips.append(ws)
        else:
            break

    # файл с описаниями прочитан, информация собрана в записи-представители класса WarShip и
    # упакована в список TheWarShips (смотреть функцию DoIt)

    # Это контрольная распечатка результатов
    for ws in WarShip.TheWarShips:
        print(ws.numberOfShip, ws.classOfShip, ws.descriptionOfShip)

    # ====================================================================

# Объект загруженного из файла/создаваемого документа .docx
#   Синтаксис:
# import docx
# doc = docx.Document(docx=None)
#   Параметры:
# docx=None - может быть либо путем к файлу .docx (строка), либо к файловым объектом.
#   Возвращаемое значение:
# объект документа Document.
#   Описание:
# Класс docx.Document() представляет собой загруженный документ, переданный docx,
# где аргумент docx может быть либо путем к файлу .docx (строка), либо к файловым объектом.
#
# Если docx=None или отсутствует, то загружается встроенный "шаблон" документа по умолчанию.
#
#
# Свойства и методы объекта Document.
#           Методы
# Document.add_heading() добавляет абзац заголовка,
# Document.add_page_break() добавляет разрыв страницы,
# Document.add_paragraph() добавляет абзац,
# Document.add_picture() добавляет изображения в отдельный абзац,
# Document.add_section() добавляет новую секцию,
# Document.add_table() добавляет новую таблицу,
# Document.save() сохраняет этот документ,
#           Свойства
# Document.core_properties основные свойства документа,
# Document.inline_shapes список объектов изображений InlineShape,
# Document.paragraphs список объектов абзацев Paragraph,
# Document.sections список объектов раздела Section,
# Document.settings объект Settings,
# Document.styles объект Styles,
# Document.tables список объектов таблиц Table,
#

    print("**********************************************************************")


    # ====================================================================
    # Document.add_heading(text='', level=1):
    # Метод Document.add_heading() добавляет абзац с текстом text и отформатированный как
    # заголовок. Абзац добавляется в конец документа. Метод возвращает ссылку на объект
    # этого абзаца.
    # Стиль заголовка будет определяться уровнем level. Если уровень равен 0, то
    # устанавливается стиль 'Title'. Если уровень 1 (или опущен), то используется заголовок
    # 'Heading 1'. Эти стили определены в интерфейсе MS Word, и если стоит ЛОКАЛИЗОВАННАЯ версия
    # Word, то эти названия стилей будут переведены на "родной" язык, например "Заголовок 1".
    # Можно не пользоваться этим методом, а добавлять заголовки вручную, используя метод
    # Document.add_paragraph() с последующим форматированием... from docx import Document
    # from docx.shared import Pt
    #
    # doc = Document()
    # # добавляется текст прогоном
    # run = doc.add_paragraph().add_run('Заголовок, размером 24 pt.')
    # # размер шрифта
    # run.font.size = Pt(24)
    # run.bold = True
    # doc.save('test.docx')
    #
    # Метод Paragraph.add_run() добавляет прогон к абзацу, содержащий текст text и имеющий
    # стиль символов style, идентифицируемый стилем идентификатора стиля.
    # Аргумент style - это встроенный или созданный пользователем стиль в интерфейсе MS Word.
    # Типа Paragraph.style, только тот, который можно применить к символам текста, а не к абзацу
    # целиком. Может быть объектом стиля абзаца Style.
    # Аргумент text может содержать символы табуляции \t, которые преобразуются в соответствующую
    # XML-форму, а также может содержать символы новой строки \n или возврата каретки \r,
    # каждый из которых преобразуется в разрыв строки.
    # Метод возвращает объект прогона Run, который дает больше контроля над форматированием
    # текста этого прогона (например: выделение цветом, другой шрифт, его размер и т.д.).
    # Прогоны в основном используются для пользовательского форматирования текста внутри абзаца.
    # Форматирование на уровне символов, например полужирный и курсив, применяется на уровне
    # прогона paragraph.add_run().
    # Прогонов в абзаце может быть больше одного.
    # Таким образом, для абзаца с полужирным словом посередине требуется три прогона:
    # нормальный, полужирный - содержащий слово, и еще один нормальный для текста после него.
    # ====================================================================

    document = Document()


    p = document.add_heading().add_run('ФЛОТ')
    p.font.size = Inches(0.25)
    p.font.italic = True
    p.font.underline = True
    p.font.color.rgb = RGBColor(100, 75, 5)

    p = document.add_paragraph().add_run('Перечень кораблей флота ')
    p.font.size = Inches(0.15)
    p.font.bold = True
    p.font.color.rgb = RGBColor(55, 25, 150)

    # пустой параграф
    p = document.add_paragraph()

    run = p.add_run('Документ составлен на основе ')
    run.font.size = Inches(0.15)
    run.font.bold = False
    run.font.color.rgb = RGBColor(250, 60, 175)


    run = p.add_run(' "анализа структуры флота" ')
    run.font.size = Inches(0.19)
    run.font.bold = True
    run.font.color.rgb = RGBColor(140, 10, 155)
    # два прогона в параграфе с разными значениями свойств
    # font.size,
    # font.bold,
    # font.color.rgb

    # ====================================================================
    # Document.add_paragraph(text='', style=None):
    # Метод Document.add_paragraph() добавляет абзац в конец документа, заполненный
    # текстом text и имеющий стиль абзаца style. Возвращает ссылку на объект добавленного
    # абзаца Paragraph.
    #
    # Аргумент text может содержать символы табуляции \t, которые преобразуются в соответствующую
    # XML-форму. Текст также может содержать символы новой строки \n или возврата каретки \r,
    # каждый из которых преобразуется в разрыв строки.
    #
    # Аргумент style может принимать строку с Именем Стиля или объект стиля Style.
    #
    # Значением аргумента style может быть Имя Стиля, которое встроено в интерфейс MS Word.
    # Встроенные стили хранятся в файле WordprocessingML под своим английским именем, например
    # 'Heading 1', и не зависят от локализации MS Word. Так как модуль python-docx работает
    # с файлом WordprocessingML, то и поиск стиля должен использовать английское имя.
    # Если файл WordprocessingML не найден (MS Word не установлен, например в OS Linux)
    # то модуль python-docx работает со своей версией этого файла.
    #
    # Пользовательские стили, которые ВЫ сами настроили, не локализованы и доступны по имени,
    # как они отображается в пользовательском интерфейсе MS Word.
    #
    # Пример добавления абзаца с текстом и стилем:
    #
    # from docx import Document
    # from docx.shared import Pt
    #
    # doc = Document()
    # # добавляем абзацы
    # doc.add_paragraph('Абзацы в Word имеют основополагающее значение.')
    # doc.add_paragraph('Стиль абзаца как цитата', style='Intense Quote')
    # doc.add_paragraph('Обычный список.', style='List Bullet')
    # doc.add_paragraph('Обычный список.', style='List Bullet')
    # # можно применить стиль и так
    # doc.add_paragraph('Нумерованный список.').style='List Number'
    # # а можно применить стиль позже, к объекту абзаца `p`
    # p = doc.add_paragraph('Нумерованный список.')
    # p.style='List Number'
    # doc.save('test.docx')

    # ====================================================================

    document.add_paragraph('Перечень кораблей', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.add_picture('DreadNought_F.png', width=Inches(5.0))

    records = []
    for warShip in WarShip.TheWarShips:
        records.append((warShip.numberOfShip, warShip.classOfShip, warShip.descriptionOfShip)),

    # records = tuple(records)


    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # таблица из одной сторки (пока из одной) - из неё
    # будет сформирован заголовок.
    # hdr_cells - вспомогательная ссылка на верхнюю (она пока одна)
    # строку таблицы. Задаются названия столбцов и ширина ячейки-заголовка.
    # По нему будет выровнены ячейки в столбце таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'n'
    hdr_cells[0].width = Inches(0.35)
    hdr_cells[1].text = 'class'
    hdr_cells[1].width = Inches(0.90)
    hdr_cells[2].text = 'description'
    hdr_cells[2].width = Inches(5.90)

    #  заполнение таблицы: строка ячеек добавляется к таблице
    #  сначала в таблице нет ничего, кроме строки-заголовка.
    #  ЗДЕСЬ (!!!) верхняя строка автоматически определяет параметры
    #  новой строки. Можно ли сделать в таблице строки с
    #  разными параметрами - пока не знаю.
    for record in records:
        row_cells = table.add_row().cells      # добавили строку к таблице
        row_cells[0].text = str(record[0])
        row_cells[1].text = record[1]
        row_cells[2].text = record[2]          # ... и записали в неё информацию из record

    # record end time
    end = time.time()

    document.add_page_break()

    # print the difference between start
    # and end time in milli. secs
    # print("The time of execution of above program is :",
    #       (end - start) * 10 ** 3, "ms")

    fullTime = str((end - start) * 10 ** 3) + " ms"

    p = document.add_paragraph()
    run = p.add_run('The time of execution with document is : ' + fullTime)
    run.font.size = Inches(0.25)
    run.font.italic = True
    run.font.bold = False
    run.font.color.rgb = RGBColor(255, 15, 75)



    document.save('mainTest.docx')

# ========================================================================

    # запись в .csv файл =================================================
    # Файл CSV (значения, разделенные запятыми) позволяет сохранять данные в табличной
    # структуре с расширением .csv. CSV-файлы используются в приложениях электронной коммерции,
    # поскольку их легко обрабатывать. Некоторые из областей, где они были использованы, включают:
    # импорт и экспорт данных клиентов
    # импорт и экспорт продукции
    # экспорт заказов
    # экспорт аналитических отчетов по электронной коммерции

    # Модули для чтения и записи
    # Модуль CSV имеет несколько функций и классов, доступных для чтения и записи CSV, и
    # они включают:
    # функция csv.reader
    # функция csv.writer
    # класс csv.Dictwriter
    # класс csv.DictReader

    csvm = universalFileMaster('csvFleetFile.csv', 'w')
    with csvm.fm:
        fieldnames = ['n', 'class', 'description']

        writer = csv.DictWriter(csvm.fm, fieldnames=fieldnames)
        writer.writeheader()

        for warShip in WarShip.TheWarShips:
            writer.writerow({'n': warShip.numberOfShip,
                             'class': warShip.classOfShip,
                             'description': warShip.descriptionOfShip})

# ========================================================================
    print('~~~~~~~~~~~~~~~~~~~json~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~')

    # список словарей: именно СПИСОК! Иначе json.load не сможет прочитать всё,
    # что было записано json.dump

    shipsList = []
    for warShip in WarShip.TheWarShips:
        ship = {'n': warShip.numberOfShip,
                'class': warShip.classOfShip,
                'description': warShip.descriptionOfShip}
        shipsList.append(ship)

    # эапись списка словарей с применением класса universalFileMaster ====
    write_file = universalFileMaster('jsonFleetFile.json', 'w')
    json.dump(shipsList, write_file.fm)

    # эапись списка словарей с применением встроенной функции ============
    # with open("jsonFleetFile.json", "w") as write_file:
    #     json.dump(shipsList, write_file)

    # чтение списка словарей с применением класса universalFileMaster ====
    read_file = universalFileMaster('jsonFleetFile.json', 'r')
    ships = json.load(read_file.fm)

    # чтение списка словарей с применением встроенной функции ============
    # with open("jsonFleetFile.json", "r") as read_file:
    #         ships = json.load(read_file)

    # преимущество класа universalFileMaster в том, что это УНИВЕРСАЛЬНЫЙ КЛАСС
    # =========================================================================

    # парсинг списка словарей ============================================
    shipsList = list(ships)
    for sh in shipsList:
        print(sh)

# ========================================================================

def main():
    DoIt()


if __name__ == "__main__":
    main()
